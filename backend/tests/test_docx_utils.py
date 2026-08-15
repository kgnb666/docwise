"""Word (.docx) 解析测试：用 python-docx 现场生成 docx 再提取，验证段落与表格。"""

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.core.docx_utils import extract_docx_text
from app.deps import get_embedder
from app.main import app
from app.rag.embedder_hash import HashEmbedder
from app.storage.document_store import reset_document_store
from app.storage.vector_store import reset_vector_store


def make_docx_bytes() -> bytes:
    """用 python-docx 生成一个含标题/段落/表格的测试 docx。"""
    doc = Document()
    doc.add_heading("DocWise 测试文档", level=1)
    doc.add_paragraph("第一段：可观测性三大支柱是日志、指标、链路追踪。")
    doc.add_paragraph("第二段：DocWise-X 协议端口是 9801。")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "阈值"
    table.rows[1].cells[0].text = "CPU"
    table.rows[1].cells[1].text = "85%"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extract_docx_paragraphs_and_table():
    text = extract_docx_text(make_docx_bytes())
    assert "DocWise 测试文档" in text
    assert "可观测性三大支柱是日志、指标、链路追踪" in text
    assert "DocWise-X 协议端口是 9801" in text
    # 表格以管道分隔输出
    assert "指标 | 阈值" in text
    assert "CPU | 85%" in text


@pytest.fixture(autouse=True)
def _isolated():
    reset_vector_store()
    reset_document_store()
    app.dependency_overrides[get_embedder] = lambda: HashEmbedder()
    yield
    app.dependency_overrides.pop(get_embedder, None)
    reset_vector_store()
    reset_document_store()


def test_upload_docx_end_to_end():
    """端到端：docx 上传 → 分块 → 入库。"""
    client = TestClient(app)
    resp = client.post(
        "/api/documents",
        files={"file": ("测试文档.docx", make_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["chunk_count"] >= 1
    assert data["name"] == "测试文档.docx"

    stats = client.get("/api/documents/stats").json()
    assert stats["documents"] == 1
